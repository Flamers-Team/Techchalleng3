"""
Gera versão DOCX do RELATORIO_TECNICO_PARA_EQUIPE.md
para distribuição em email/teams.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

SRC = Path(r"C:\Users\Teste\Downloads\Techchalleng3\docs\RELATORIO_TECNICO_PARA_EQUIPE.md")
OUT = Path(r"C:\Users\Teste\Downloads\RELATORIO_TECNICO_TECHCHALLENGE_FASE3.docx")

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_md(text):
    """Adiciona parágrafo com formatação markdown básica."""
    if text.startswith('# '):
        p = doc.add_heading(text[2:], level=1)
    elif text.startswith('## '):
        p = doc.add_heading(text[3:], level=2)
    elif text.startswith('### '):
        p = doc.add_heading(text[4:], level=3)
    elif text.startswith('#### '):
        p = doc.add_heading(text[5:], level=4)
    elif text.startswith('|'):
        # Tabela markdown - pula (simplificado)
        return None
    elif text.startswith('- '):
        return doc.add_paragraph(text[2:], style='List Bullet')
    elif text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
        return doc.add_paragraph(text, style='List Number')
    elif text.startswith('```'):
        return None  # bloco de código, pula
    elif text.startswith('---'):
        # separador
        return doc.add_paragraph('─' * 50)
    elif text.strip() == '':
        return None
    else:
        # Negrito inline
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                # Inline code
                code_parts = re.split(r'(`[^`]+`)', part)
                for cp in code_parts:
                    if cp.startswith('`') and cp.endswith('`'):
                        run = p.add_run(cp[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                    else:
                        p.add_run(cp)
        return p


# Ler markdown e adicionar
with SRC.open(encoding='utf-8') as f:
    for line in f:
        line = line.rstrip('\n')
        if not line.strip():
            continue
        # Ignorar tabelas markdown (complexas demais)
        if line.startswith('|') and '---' not in line:
            continue
        if line.startswith('|---'):
            continue
        add_md(line)

doc.save(OUT)
print(f"✅ DOCX gerado: {OUT}")
print(f"   Tamanho: {OUT.stat().st_size / 1024:.1f} KB")