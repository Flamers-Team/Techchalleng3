"""
Conversor robusto de Markdown para DOCX com tabelas.
Suporta: cabeçalhos, listas, tabelas, code blocks, bold, italic, code inline.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys


def set_cell_shading(cell, color_hex: str):
    """Aplica cor de fundo a uma célula de tabela."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_formatted_run(paragraph, text: str, base_bold=False, base_italic=False, base_code=False):
    """Adiciona texto com formatação inline (**bold**, *italic*, `code`)."""
    # Padrão para **bold**, *italic*, `code`
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"

    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            # Markdown link [text](url)
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
                run.underline = True
        else:
            paragraph.add_run(part)


def md_to_docx(md_text: str, output_path: str):
    doc = Document()

    # Estilo padrão
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Code block delimiters
        if line.strip().startswith("```"):
            if in_code_block:
                # Fecha code block
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Cabeçalhos
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif line.startswith("#### "):
            p = doc.add_heading(line[5:].strip(), level=4)

        # Tabelas (formato markdown)
        elif line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i+1].strip()):
            # Cabeçalho da tabela
            header_cells = [c.strip() for c in line.strip("|").split("|")]
            table = doc.add_table(rows=1, cols=len(header_cells))
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for j, cell_text in enumerate(header_cells):
                hdr[j].text = ""
                p = hdr[j].paragraphs[0]
                run = p.add_run(cell_text)
                run.bold = True
                run.font.size = Pt(10)
                set_cell_shading(hdr[j], "1F497D")
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            # Pula linha separadora
            i += 2

            # Linhas de dados
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip("|").split("|")]
                row = table.add_row().cells
                for j, cell_text in enumerate(cells):
                    if j < len(row):
                        row[j].text = ""
                        p = row[j].paragraphs[0]
                        add_formatted_run(p, cell_text)
                        for run in p.runs:
                            run.font.size = Pt(9)
                i += 1
            continue  # já avançamos i

        # Linha de separação (já tratada acima)
        elif re.match(r"^\|[\s\-:|]+\|$", line.strip()):
            i += 1
            continue

        # Listas com bullet
        elif re.match(r"^\s*[-*]\s+", line):
            content = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_run(p, content)

        # Listas numeradas
        elif re.match(r"^\s*\d+\.\s+", line):
            content = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_formatted_run(p, content)

        # Blockquote
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            run = p.add_run("┃ " + line[2:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Linha horizontal
        elif re.match(r"^---+$", line.strip()):
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # Linha vazia
        elif not line.strip():
            pass  # Ignora linhas vazias consecutivas

        # Texto normal
        else:
            p = doc.add_paragraph()
            add_formatted_run(p, line)

        i += 1

    doc.save(output_path)
    print(f"OK: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Uso: python md_to_docx.py input.md output.docx")
        sys.exit(1)
    md_path = sys.argv[1]
    docx_path = sys.argv[2]
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()
    md_to_docx(md_text, docx_path)
